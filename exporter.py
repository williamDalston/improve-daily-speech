"""
Export functions for .txt, .docx, and audio (OpenAI TTS).
"""

import io
import os

import openai
from docx import Document
from pydub import AudioSegment
from docx.shared import Pt
from dotenv import load_dotenv

load_dotenv()


def _get_openai_client():
    """Lazy client init — works on Streamlit Cloud where secrets aren't in env."""
    try:
        import streamlit as st
        if "OPENAI_API_KEY" not in os.environ and "OPENAI_API_KEY" in st.secrets:
            os.environ["OPENAI_API_KEY"] = st.secrets["OPENAI_API_KEY"]
    except Exception:
        pass
    return openai.OpenAI()


def export_txt(text: str) -> str:
    return text


def export_docx(text: str, title: str) -> bytes:
    doc = Document()
    doc.add_heading(title, level=1)

    paragraphs = text.split("\n\n")
    for para_text in paragraphs:
        stripped = para_text.strip()
        if stripped:
            p = doc.add_paragraph(stripped)
            for run in p.runs:
                run.font.size = Pt(12)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generate_audio(text: str, voice: str = "onyx", speed: float = 1.0) -> bytes:
    """
    Generate documentary-style audio from text using OpenAI TTS.

    OpenAI TTS has a 4096 character limit per request, so we split
    long texts into chunks and properly merge the audio using pydub
    to avoid glitches at chunk boundaries.

    Voices: alloy, ash, ballad, coral, echo, fable, onyx, nova, sage, shimmer
    Speed: 0.25 to 4.0 (1.0 = normal)

    Returns MP3 bytes.
    """
    chunks = _split_for_tts(text, max_chars=4000)

    # Generate audio for each chunk
    audio_segments = []
    for chunk in chunks:
        response = _get_openai_client().audio.speech.create(
            model="tts-1-hd",
            voice=voice,
            input=chunk,
            speed=speed,
            response_format="mp3",
        )
        # Load MP3 bytes into pydub AudioSegment
        segment = AudioSegment.from_mp3(io.BytesIO(response.content))
        audio_segments.append(segment)

    # Merge segments with small crossfade to avoid glitches
    if len(audio_segments) == 1:
        combined = audio_segments[0]
    else:
        combined = audio_segments[0]
        for segment in audio_segments[1:]:
            # 50ms crossfade for seamless joins
            combined = combined.append(segment, crossfade=50)

    # Export as MP3
    buffer = io.BytesIO()
    combined.export(buffer, format="mp3", bitrate="192k")
    buffer.seek(0)
    return buffer.getvalue()


def _split_for_tts(text: str, max_chars: int = 4000) -> list[str]:
    """Split text into chunks at paragraph boundaries, staying under max_chars."""
    paragraphs = text.split("\n\n")
    chunks = []
    current = ""

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue

        # If a single paragraph exceeds max, split at sentence boundaries
        if len(para) > max_chars:
            if current:
                chunks.append(current.strip())
                current = ""
            sentences = _split_sentences(para)
            for sentence in sentences:
                if len(current) + len(sentence) + 1 > max_chars:
                    if current:
                        chunks.append(current.strip())
                    current = sentence
                else:
                    current = current + " " + sentence if current else sentence
        elif len(current) + len(para) + 2 > max_chars:
            chunks.append(current.strip())
            current = para
        else:
            current = current + "\n\n" + para if current else para

    if current.strip():
        chunks.append(current.strip())

    return chunks


def _split_sentences(text: str) -> list[str]:
    """Rough sentence splitting on . ! ?"""
    import re
    parts = re.split(r'(?<=[.!?])\s+', text)
    return [p for p in parts if p.strip()]
